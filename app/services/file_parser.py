"""
文件解析服务。

支持从 .docx 等格式中提取纯文本内容。
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str | Path) -> str:
    """
    从 .docx 文件中提取纯文本。

    Args:
        file_path: .docx 文件路径

    Returns:
        提取的纯文本内容

    Raises:
        ValueError: 解析失败时抛出
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    try:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"无法解析 .docx 文件: {e}")


def extract_text_from_docx_bytes(data: bytes) -> str:
    """
    从 .docx 二进制数据中提取纯文本。

    Args:
        data: .docx 文件的二进制内容

    Returns:
        提取的纯文本内容
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    try:
        import io
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"无法解析 .docx 数据: {e}")
